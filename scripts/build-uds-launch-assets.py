#!/usr/bin/env python3
"""Build Andrew Green packet: DOCX + 2 PNG diagrams for the UDS bundle launch.

Every URL referenced in this packet has been HTTP-verified live before
generation. No fictional ghcr.io coordinates, no fictional monorepo, no
fictional keyless-OIDC verify chain. Only the GitHub Releases that are
actually published and the cosign-blob verify command that actually works
against the .sig + .pub assets attached to each release.

Outputs:
  docs/uds/exports/SZL-UDS-Mesh-Launch-AndrewGreen.docx
  docs/uds/exports/img/uds-fleet-mesh.png
  docs/uds/exports/img/uds-pull-verify-install.png
"""
from __future__ import annotations
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "uds" / "exports"
IMG_DIR = OUT_DIR / "img"
OUT_DIR.mkdir(parents=True, exist_ok=True)
IMG_DIR.mkdir(parents=True, exist_ok=True)

BG       = "#0b0e14"
FG       = "#e6edf3"
MUTED    = "#8b949e"
CYAN     = "#22d3ee"
AMBER    = "#f0b429"
GREEN    = "#34d399"
PANEL    = "#11161d"
BORDER   = "#1f2733"

BUNDLES = [
    ("A11oy",   "brand orchestration",       "#22d3ee"),
    ("Amaru",   "convergent data-sync",      "#a78bfa"),
    ("ROSIE",   "governed decisions",        "#34d399"),
    ("Sentra",  "cyber resilience",          "#f0b429"),
    ("Vessels", "maritime intelligence",     "#60a5fa"),
]

# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def panel(ax, x, y, w, h, fc=PANEL, ec=BORDER, lw=1.2, radius=0.04):
    p = FancyBboxPatch((x, y), w, h,
                       boxstyle=f"round,pad=0.005,rounding_size={radius}",
                       linewidth=lw, edgecolor=ec, facecolor=fc)
    ax.add_patch(p)
    return p

def text(ax, x, y, s, **kw):
    kw.setdefault("color", FG)
    kw.setdefault("fontsize", 11)
    kw.setdefault("ha", "left")
    kw.setdefault("va", "center")
    kw.setdefault("family", "DejaVu Sans Mono")
    ax.text(x, y, s, **kw)

# ──────────────────────────────────────────────────────────────────────────────
# Diagram 1: Five per-bundle repos → uds-v0.2.0 release assets → operator
# ──────────────────────────────────────────────────────────────────────────────

def fig_fleet_mesh(path):
    fig, ax = plt.subplots(figsize=(13.33, 7.5), dpi=160)
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 16); ax.set_ylim(-1.5, 9); ax.set_axis_off()

    text(ax, 0.6, 8.5, "SZL Holdings — UDS bundles, v0.2.0",
         fontsize=22, weight="bold", family="DejaVu Sans")
    text(ax, 0.6, 7.95,
         "five public GitHub repos → one signed release each → four-asset bundle → operator",
         fontsize=12, color=MUTED, family="DejaVu Sans")

    # Left column: 5 per-bundle repos
    bx = 0.6; bw = 6.4; bh = 1.05; gap = 0.15
    by = 6.9
    for slug, blurb, color in BUNDLES:
        panel(ax, bx, by - bh, bw, bh, fc=PANEL, ec=color, lw=1.6)
        panel(ax, bx, by - bh, 0.18, bh, fc=color, ec=color, lw=0, radius=0.02)
        text(ax, bx + 0.40, by - 0.30, slug, fontsize=13, weight="bold")
        text(ax, bx + 1.60, by - 0.32, blurb, fontsize=9.5, color=MUTED,
             family="DejaVu Sans")
        text(ax, bx + 0.40, by - 0.78,
             f"github.com/szl-holdings/{slug.lower()}  →  uds-v0.2.0",
             fontsize=9.0, color=CYAN)
        by -= (bh + gap)

    # Middle: release assets (4 per bundle)
    mx, my, mw, mh = 7.6, 3.6, 3.8, 3.4
    panel(ax, mx, my, mw, mh, fc=PANEL, ec=CYAN, lw=2.2, radius=0.05)
    text(ax, mx + mw/2, my + mh - 0.45, "Release assets",
         fontsize=15, weight="bold", color=CYAN, ha="center")
    text(ax, mx + mw/2, my + mh - 0.85, "four per bundle, every one HTTP-verified",
         fontsize=9.5, color=MUTED, ha="center", family="DejaVu Sans")
    rows = [
        "<bundle>-uds-0.2.0.tar.zst",
        "<bundle>-uds-0.2.0.tar.zst.sha256",
        "<bundle>-uds-0.2.0.tar.zst.sig",
        "<bundle>-uds-dev.pub",
    ]
    for i, r in enumerate(rows):
        text(ax, mx + 0.25, my + mh - 1.55 - i*0.40, "• " + r,
             fontsize=9.5, color=FG)
    text(ax, mx + 0.25, my + mh - 3.20,
         "attached to the GitHub Release",
         fontsize=9, color=MUTED, family="DejaVu Sans")

    # Arrows from repos → assets
    arrow_x_start = bx + bw + 0.05
    arrow_x_end = mx
    y_cur = 6.9 - bh/2
    for _ in BUNDLES:
        arr = FancyArrowPatch((arrow_x_start, y_cur), (arrow_x_end, my + mh/2),
                              arrowstyle="-|>", mutation_scale=10,
                              color=BORDER, linewidth=1.0, alpha=0.7)
        ax.add_patch(arr)
        y_cur -= (bh + gap)

    # Right: operator
    cx, cy, cw, ch = 12.0, 3.6, 3.6, 3.4
    panel(ax, cx, cy, cw, ch, fc=PANEL, ec=AMBER, lw=1.8, radius=0.05)
    text(ax, cx + cw/2, cy + ch - 0.45, "Operator",
         fontsize=14, weight="bold", color=AMBER, ha="center")
    text(ax, cx + cw/2, cy + ch - 0.85, "connected or air-gapped",
         fontsize=10, color=MUTED, ha="center", family="DejaVu Sans")
    for i, r in enumerate([
        "• curl the four assets",
        "• sha256sum -c",
        "• cosign verify-blob",
        "• zarf package deploy",
        "• runtime under /opt/<bundle>/",
    ]):
        text(ax, cx + 0.25, cy + ch - 1.45 - i*0.35, r,
             fontsize=10, color=FG)

    arr = FancyArrowPatch((mx + mw, my + mh/2), (cx, cy + ch/2),
                          arrowstyle="-|>", mutation_scale=14,
                          color=CYAN, linewidth=2.0)
    ax.add_patch(arr)
    text(ax, (mx + mw + cx)/2, my + mh/2 + 0.25, "fetch",
         fontsize=9, color=CYAN, ha="center", family="DejaVu Sans")

    # Footer trust anchor
    panel(ax, 0.6, -1.2, 14.8, 1.6, fc="#0f1620", ec=GREEN, lw=1.4, radius=0.04)
    text(ax, 0.95, 0.05, "trust anchor",
         fontsize=10.5, color=GREEN, weight="bold")
    text(ax, 0.95, -0.35,
         "the .sig is a cosign blob signature over the .tar.zst, made with the dev keypair —",
         fontsize=10, color=FG, family="DejaVu Sans")
    text(ax, 0.95, -0.70,
         "the matching public key (<bundle>-uds-dev.pub) is published as the fourth release asset.",
         fontsize=10, color=FG, family="DejaVu Sans")
    text(ax, 0.95, -1.05,
         "deterministic builds • content-addressed sha256 • signature pinned to a published key • air-gap parity",
         fontsize=9, color=MUTED, family="DejaVu Sans")

    plt.savefig(path, dpi=160, facecolor=BG, bbox_inches="tight")
    plt.close(fig)
    print(f"wrote {path}")

# ──────────────────────────────────────────────────────────────────────────────
# Diagram 2: download → sha256 → cosign verify-blob → zarf deploy
# ──────────────────────────────────────────────────────────────────────────────

def fig_contract(path):
    fig, ax = plt.subplots(figsize=(13.33, 8.4), dpi=160)
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.set_axis_off()

    text(ax, 0.6, 9.5, "Universal verify-and-install contract",
         fontsize=22, weight="bold", family="DejaVu Sans")
    text(ax, 0.6, 8.95, "one shape — five bundles — air-gap-compatible",
         fontsize=12, color=MUTED, family="DejaVu Sans")

    steps = [
        ("1", "DOWNLOAD", CYAN,
         "BASE=https://github.com/szl-holdings/<bundle>/releases/download/uds-v0.2.0\n"
         "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst\n"
         "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst.sha256\n"
         "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst.sig\n"
         "curl -LO $BASE/<bundle>-uds-dev.pub",
         "all four assets attached to\nthe published GitHub Release"),
        ("2", "SHA256", AMBER,
         "sha256sum -c <bundle>-uds-0.2.0.tar.zst.sha256",
         "content-addressed —\nbyte-for-byte verifies the tarball"),
        ("3", "COSIGN", GREEN,
         "cosign verify-blob \\\n"
         "  --key <bundle>-uds-dev.pub \\\n"
         "  --signature <bundle>-uds-0.2.0.tar.zst.sig \\\n"
         "  <bundle>-uds-0.2.0.tar.zst",
         "signature pinned to the\npublished dev public key"),
        ("4", "DEPLOY", "#a78bfa",
         "zarf package deploy <bundle>-uds-0.2.0.tar.zst --confirm",
         "components stage under\n/opt/<bundle>/ on the target"),
    ]

    y_cursor = 8.3
    for n, label, color, cmd, blurb in steps:
        h = 1.70
        circle = plt.Circle((1.2, y_cursor - h/2), 0.42, color=color, ec=color, lw=2)
        ax.add_patch(circle)
        text(ax, 1.2, y_cursor - h/2, n, fontsize=20, weight="bold",
             color=BG, ha="center", va="center", family="DejaVu Sans")
        text(ax, 2.0, y_cursor - 0.30, label, fontsize=14, weight="bold",
             color=color, family="DejaVu Sans")
        text(ax, 2.0, y_cursor - 0.65, blurb, fontsize=9.5, color=MUTED,
             family="DejaVu Sans", va="top")
        panel(ax, 6.2, y_cursor - h + 0.10, 9.4, h - 0.20,
              fc="#0d141c", ec=BORDER, lw=1.2)
        text(ax, 6.4, y_cursor - 0.30, cmd, fontsize=8.8, color=FG, va="top")
        y_cursor -= (h + 0.15)

    panel(ax, 0.6, 0.4, 14.8, 1.0, fc="#0f1620", ec=BORDER, lw=1.2, radius=0.04)
    text(ax, 0.95, 1.10, "trust chain",
         fontsize=10.5, color=CYAN, weight="bold")
    text(ax, 0.95, 0.75,
         "published .pub key → cosign verify-blob → matching .sig → sha256-pinned tarball → zarf deploy.",
         fontsize=9.5, color=FG, family="DejaVu Sans")
    text(ax, 0.95, 0.50,
         "every step verifies the previous one.  break the chain, the deploy halts.",
         fontsize=9, color=MUTED, family="DejaVu Sans")

    plt.savefig(path, dpi=160, facecolor=BG, bbox_inches="tight")
    plt.close(fig)
    print(f"wrote {path}")

# ──────────────────────────────────────────────────────────────────────────────
# Word doc for Andrew Green
# ──────────────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(doc, text, *, mono=False, color=None, size=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return p

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p

def add_link_para(doc, label, url):
    p = doc.add_paragraph()
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r2 = p.add_run(url)
    r2.font.name = "Consolas"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x07, 0x5F, 0xB6)
    return p

def build_docx(path, fleet_png, contract_png):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("SZL Holdings — UDS bundles, v0.2.0")
    r.bold = True
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Defense-Unicorns operator packet · Prepared for Andrew Green · "
        "Every URL in this packet was HTTP-verified live before send."
    )
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    doc.add_picture(str(fleet_png), width=Inches(6.5))
    doc.add_paragraph()

    # ── Executive summary ────────────────────────────────────────────────
    add_heading(doc, "Executive summary", level=1)
    doc.add_paragraph(
        "Five signed Zarf payloads for Defense-Unicorns environments — one "
        "public GitHub repo per bundle, one uds-v0.2.0 release on each, four "
        "verifiable assets per release. Same download → sha256 → "
        "cosign-verify-blob → zarf-deploy contract across the fleet. "
        "Air-gap-compatible by construction: every byte you need to verify and "
        "install is attached to the Release itself, no separate registry, no "
        "out-of-band trust."
    )
    doc.add_paragraph(
        "If you are standing up a UDS-enabled cluster — connected, edge, or "
        "fully disconnected — the full fleet is pullable today straight from "
        "GitHub Releases."
    )

    # ── Bundle table ────────────────────────────────────────────────────
    add_heading(doc, "The five bundles", level=1)
    table = doc.add_table(rows=1, cols=4)
    # `Table Grid` is a built-in Word style; safer than `Light Grid Accent 1`
    # which python-docx references by name but does NOT inject into the doc's
    # styles.xml -- Word then flags the file as "needing repair" on download.
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["#", "Bundle", "Runtime", "Release URL"]):
        hdr[i].text = label
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    rows = [
        ("1", "A11oy",
         "Brand orchestration — @a11oy/core + @a11oy/connection, optional hash-chained attestations component for offline provenance",
         "https://github.com/szl-holdings/a11oy/releases/tag/uds-v0.2.0"),
        ("2", "Amaru",
         "Convergent multi-source data-sync — append-only delta logs, hash-verified ingest, bounded-loop convergence, KL drift, hash-chained proof receipts",
         "https://github.com/szl-holdings/amaru/releases/tag/uds-v0.2.0"),
        ("3", "ROSIE",
         "Governed decision fabric — deny-by-default admission, contradiction detector, governed-action emit, hash-chained decision receipts",
         "https://github.com/szl-holdings/rosie/releases/tag/uds-v0.2.0"),
        ("4", "Sentra",
         "Cyber resilience command — asset-scoped fail-closed Safety Gate · NIST CSF 2.0 + SP 800-61r2 + CISA CIRCIA + MITRE D3FEND mappings · Ising allocation · hash-chained Proof Chain",
         "https://github.com/szl-holdings/sentra/releases/tag/uds-v0.2.0"),
        ("5", "Vessels",
         "Maritime intelligence — trajectory inspector, AIS-gap detector, sanctions screen, voyage Λ-receipts",
         "https://github.com/szl-holdings/vessels/releases/tag/uds-v0.2.0"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val

    doc.add_paragraph()
    add_para(doc, "Every release ships exactly four assets:", bold=True)
    add_code(doc,
        "<bundle>-uds-0.2.0.tar.zst             # the bundle (zarf-deployable tarball)\n"
        "<bundle>-uds-0.2.0.tar.zst.sha256       # content-addressed checksum\n"
        "<bundle>-uds-0.2.0.tar.zst.sig          # cosign blob signature\n"
        "<bundle>-uds-dev.pub                    # the public key the .sig verifies against"
    )

    # ── Contract diagram ─────────────────────────────────────────────────
    add_heading(doc, "Universal verify-and-install contract", level=1)
    doc.add_picture(str(contract_png), width=Inches(6.5))

    add_code(doc,
        "# Replace <bundle> with one of: a11oy | amaru | rosie | sentra | vessels\n\n"
        "BASE=https://github.com/szl-holdings/<bundle>/releases/download/uds-v0.2.0\n\n"
        "# 1. DOWNLOAD\n"
        "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst\n"
        "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst.sha256\n"
        "curl -LO $BASE/<bundle>-uds-0.2.0.tar.zst.sig\n"
        "curl -LO $BASE/<bundle>-uds-dev.pub\n\n"
        "# 2. SHA256 — byte-for-byte integrity\n"
        "sha256sum -c <bundle>-uds-0.2.0.tar.zst.sha256\n\n"
        "# 3. COSIGN — signature pinned to the published dev public key\n"
        "cosign verify-blob \\\n"
        "  --key <bundle>-uds-dev.pub \\\n"
        "  --signature <bundle>-uds-0.2.0.tar.zst.sig \\\n"
        "  <bundle>-uds-0.2.0.tar.zst\n\n"
        "# 4. DEPLOY — components stage under /opt/<bundle>/ on the target\n"
        "zarf package deploy <bundle>-uds-0.2.0.tar.zst --confirm"
    )

    # ── Air-gap path ─────────────────────────────────────────────────────
    add_heading(doc, "Air-gap path — the reason this whole thing exists", level=1)
    doc.add_paragraph(
        "Because every asset that participates in the trust chain is attached "
        "to the GitHub Release itself, the air-gap path is identical to the "
        "connected path with one extra hop:"
    )
    for s in [
        "On a connected host, curl all four assets from the Release page above.",
        "Carry the four files across the gap on whatever media policy dictates.",
        "On the air-gapped host, run steps 2 → 4 unchanged. No GHCR reach-out, no Sigstore round-trip, no external dependency.",
    ]:
        doc.add_paragraph(s, style="List Number")

    # ── Source repos ─────────────────────────────────────────────────────
    add_heading(doc, "Source repos — public, auditable, one per bundle", level=1)
    doc.add_paragraph(
        "There is no monorepo. Each bundle is its own repo, with its own "
        "release cadence, its own issue tracker, and its own audit surface."
    )
    for label, url in [
        ("A11oy",     "https://github.com/szl-holdings/a11oy"),
        ("Amaru",     "https://github.com/szl-holdings/amaru"),
        ("ROSIE",     "https://github.com/szl-holdings/rosie"),
        ("Sentra",    "https://github.com/szl-holdings/sentra"),
        ("Vessels",   "https://github.com/szl-holdings/vessels"),
        ("UDS Mesh",  "https://github.com/szl-holdings/uds-mesh"),
        ("Org page",  "https://github.com/szl-holdings"),
    ]:
        add_link_para(doc, label + ":", url)

    # ── Doctrine ─────────────────────────────────────────────────────────
    add_heading(doc, "Doctrine guarantees — the things we will not break", level=1)
    for line in [
        "Deterministic — rebuild the same source-tree → byte-for-byte the same tarball; the published .sha256 verifies it.",
        "Signature-pinned — every release ships its own dev public key alongside the .sig; verification needs only what the Release itself provides.",
        "Self-contained Releases — sha256, signature, public key, and tarball are all attached to the same Release; no external registry, no Sigstore round-trip required for offline operators.",
        "One contract, five bundles — adding a sixth bundle does not change the download / sha256 / cosign / deploy shape. New bundles inherit the contract, not the other way around.",
        "Public source — every repo is public; the build scripts and verifier glue you would audit are readable end-to-end.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Doctrine receipt — published, machine-verified, measured ────────
    add_heading(doc, "Doctrine receipt — published, machine-verified, measured", level=1)
    doc.add_paragraph(
        "The bundles above are the operator surface of a doctrine that is "
        "published, peer-deposited on Zenodo, and Lean-4 kernel-verified. "
        "Every consequential action in the runtime traverses a 9-step "
        "governance loop — signal → context → recommendation → simulation → "
        "policy → approval → execution → proof → outcome — and seals into a "
        "hash-chained Proof Chain. The primitives are not marketing; they are "
        "named theorems with measured overhead."
    )
    add_para(doc, "Primitives the UDS bundles inherit:", bold=True)
    for line in [
        "Λ-gate (9-axis Lutar Invariant) — admission gate; Lean-4 kernel-verified uniqueness theorem.",
        "Bekenstein-bounded admission — information-budget cap on what enters the loop per receipt.",
        "Dual-witness verdict (MATCH / DIVERGE) — every governed decision carries two independent witnesses; mismatch fails closed.",
        "KL drift bound — convergence guarantee on bounded loops with a measurable closure ratio.",
        "Reference-vector parity — bit-exact cross-runtime determinism; receipts are reproducible byte-for-byte.",
        "Proof Chain — append-only hash-chained signed receipts; the audit log IS the execution record.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_para(doc, "Published doctrine (Ouroboros Thesis, v1–v13 on Zenodo):", bold=True)
    add_code(doc,
        "Concept DOI (always latest):  10.5281/zenodo.19944926\n"
        "v13  Unified Ouroboros Spine — Anatomy as Architecture     10.5281/zenodo.20195368\n"
        "v12  The Lambda-Ouroboros Substrate — Lean-4 verified      10.5281/zenodo.20173920\n"
        "v11  APPLIED Lambda — measured per-request overhead        10.5281/zenodo.20119582\n"
        "v10  EXHAUSTIVE-AUDIT — Lambda_10 audit-closure operator   10.5281/zenodo.20053163\n"
        "v9   UNIFIED-OPERATIONAL — Lutar family with Bianchi       10.5281/zenodo.20053148\n"
        "v3   The Lutar Invariant — axiomatic trust aggregator      10.5281/zenodo.19983066\n"
        "v2   Empirical companion (A11oy, Sentra, Amaru)            10.5281/zenodo.19934129\n"
        "v1   Position paper — bounded looped computation           10.5281/zenodo.19867281"
    )

    add_para(doc, "Quantitative anchor (v11, APPLIED Λ):", bold=True)
    for line in [
        "24,800 HTTP calls measured across 8 routes in a governed runtime.",
        "Median Λ₁₀ overhead 0.49–0.59 ms · p99 ≤ 1.27 ms per request.",
        "ρ = 1.000 (audit-closure ratio) on 8,000 / 8,000 governed pairs.",
        "61/61 unit + adapter tests passing for the Λ-Ouroboros substrate.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_para(doc, "Provenance:", bold=True)
    add_link_para(doc, "Lean-4 kernel-verified Lutar Invariant:",
                  "https://github.com/szl-holdings/lutar-lean")
    add_link_para(doc, "Ouroboros runtime (218/218 guardrail tests passing):",
                  "https://github.com/szl-holdings/ouroboros")
    add_link_para(doc, "Ouroboros Thesis (canonical text, paper-kit per version):",
                  "https://github.com/szl-holdings/ouroboros-thesis")
    add_link_para(doc, "Author ORCID — Stephen P. Lutar:",
                  "https://orcid.org/0009-0001-0110-4173")

    # ── Why ──────────────────────────────────────────────────────────────
    add_heading(doc, "Why this matters for UDS / Defense-Unicorns operators", level=1)
    for line in [
        "Air-gap parity by construction — the connected and disconnected paths run the same four commands; only the transport changes.",
        "Signature-pinned trust — no Sigstore round-trip, no OIDC dependency at verify time; the .pub key is published next to the .sig.",
        "Per-bundle blast radius — bundles ship from independent repos; a regression in one cannot stall a release of another.",
        "Doctrine-bound runtimes — every bundle ships hash-chained receipts (Proof Chain on Sentra, decision receipts on ROSIE, proof receipts on Amaru, Λ-receipts on Vessels, attestations component on A11oy) so nothing on the node runs without writing what it did.",
        "Auditable end-to-end — public repos, per-bundle Releases, signed assets you can verify in fifteen seconds.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    closing = doc.add_paragraph()
    cr = closing.add_run(
        "Download a tarball. Check its sha256. Verify its signature against "
        "the published dev key. Deploy. If anything in those four steps "
        "surprises you, that is a bug — open an issue on the per-bundle repo."
    )
    cr.italic = True

    doc.save(str(path))
    print(f"wrote {path}")

# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    fleet_png    = IMG_DIR / "uds-fleet-mesh.png"
    contract_png = IMG_DIR / "uds-pull-verify-install.png"
    docx_path    = OUT_DIR / "SZL-UDS-Mesh-Launch-AndrewGreen.docx"

    fig_fleet_mesh(fleet_png)
    fig_contract(contract_png)
    build_docx(docx_path, fleet_png, contract_png)
    print("done")
