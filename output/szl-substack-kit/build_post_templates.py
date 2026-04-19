#!/usr/bin/env python3
"""Generate SZL-Substack-Post-Templates.docx — 5 ready-to-publish templates."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "/home/runner/workspace/output/szl-substack-kit/SZL-Substack-Post-Templates.docx"

GOLD = RGBColor(0xD4, 0xA0, 0x54)
BLUE = RGBColor(0x3D, 0x6F, 0xD9)
DARK_GREY = RGBColor(0x22, 0x2B, 0x38)
MUTED = RGBColor(0x64, 0x74, 0x88)
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_heading(doc, text, level=1, color=GOLD):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = GOLD
        elif level == 2:
            run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = BLUE
        elif level == 3:
            run.font.size = Pt(12); run.font.bold = True; run.font.color.rgb = DARK_GREY
    return h


def add_label_value(doc, label, value, note=None):
    p = doc.add_paragraph()
    rl = p.add_run(f"{label}: "); rl.bold = True
    rl.font.color.rgb = DARK_GREY; rl.font.size = Pt(11)
    rv = p.add_run(value); rv.font.color.rgb = BLACK; rv.font.size = Pt(11)
    if note:
        rn = p.add_run(f"  ({note})")
        rn.font.color.rgb = MUTED; rn.font.size = Pt(10); rn.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GREY
        if italic:
            run.font.italic = True
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_placeholder(doc, text):
    """Bracketed guidance line — italic, muted, indented."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    r.italic = True
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_template_block(doc, text):
    """Verbatim post body (copy-paste content)."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Note: {text}")
    r.font.size = Pt(10); r.font.color.rgb = MUTED; r.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D4A054')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_section_header(doc, label):
    """Small inline label e.g. 'HOOK', 'BODY', 'CTA'."""
    p = doc.add_paragraph()
    r = p.add_run(label.upper())
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def build_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Cover ─────────────────────────────────────────────────────────────────
    title = doc.add_heading('SZL Holdings', 0)
    for run in title.runs:
        run.font.color.rgb = GOLD; run.font.size = Pt(28); run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Substack Post Templates — 5 Formats Ready to Publish')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = MUTED; run.font.size = Pt(12); run.font.italic = True
    doc.add_paragraph()

    add_body(doc,
        "Five reusable post templates tuned to the SZL Holdings voice. Each one is designed "
        "to be filled in and published in under 60 minutes. Together they form a publishing "
        "rhythm: one weekly digest, one essay, one build log, one field report per month, and "
        "one reading list per month. Hit publish on something every week without staring at a "
        "blank page.", italic=True)
    add_divider(doc)

    # ── Voice & Brand Notes ───────────────────────────────────────────────────
    add_heading(doc, 'Voice & Brand Notes', level=1)
    add_body(doc,
        "Every template below is written in SZL's house voice. Stay inside these rails when "
        "filling in the placeholders.")
    add_heading(doc, 'Voice', level=3)
    add_body(doc, "• Builder-operator. First-person. Plain English. No jargon for jargon's sake.")
    add_body(doc, "• Proof over pitch. Lead with what shipped, what broke, or what was decided — never with marketing language.")
    add_body(doc, "• Short sentences. Active verbs. Cut every adjective that isn't earning its place.")
    add_body(doc, "• Confident, not promotional. Show the work; let the reader draw the conclusion.")
    add_body(doc, "• When in doubt: replace 'we are excited to announce' with 'here is what shipped'.")

    add_heading(doc, 'Banned phrases', level=3)
    add_body(doc, "AI-powered · revolutionary · cutting-edge · best-in-class · seamless · synergies · "
                  "leverage (as a verb) · we are excited to announce · in today's fast-paced world.")

    add_heading(doc, 'Formatting', level=3)
    add_body(doc, "• Headline: sentence case, no clickbait, ≤ 70 characters.")
    add_body(doc, "• Subhead (Substack subtitle field): one sentence that earns the click.")
    add_body(doc, "• Open with one line that lands. No throat-clearing.")
    add_body(doc, "• Use H2 (## in Substack) for major sections, H3 (###) sparingly.")
    add_body(doc, "• Bulleted lists for ≥ 3 parallel items only. Otherwise prose.")
    add_body(doc, "• Bold one phrase per section maximum. Italics for product names on first use.")
    add_body(doc, "• Sign every post: '— Stephen Lutar, Founder & CEO, SZL Holdings'.")

    add_heading(doc, 'Substack post settings (apply to every template)', level=3)
    add_label_value(doc, 'Access', 'Free — visible to everyone',
        note='Switch to paid only when premium tier is launched')
    add_label_value(doc, 'Section', 'Default')
    add_label_value(doc, 'Post type', 'Newsletter (email + web)')
    add_label_value(doc, 'Comments', 'On — all subscribers')
    add_label_value(doc, 'Cover image', 'Optional — only if it adds signal, never as decoration')
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # TEMPLATE 1 — Weekly Intelligence Digest
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, 'Template 1 — The Signal (Weekly Intelligence Digest)', level=1)
    add_body(doc,
        "The anchor post. Publishes every Monday. Short, scannable, signal-dense. "
        "Five items: one from each portfolio domain plus one cross-cutting observation. "
        "Trains readers to expect the inbox at the same time every week.", italic=True)

    add_label_value(doc, 'Cadence', 'Weekly — Monday 07:00 ET')
    add_label_value(doc, 'Target length', '450–650 words / 3–4 minute read')
    add_label_value(doc, 'Post tag', 'the-signal')

    add_heading(doc, 'Headline formula', level=2)
    add_body(doc,
        "The Signal · [Issue number] — [3-to-6-word characterisation of the week]")
    add_placeholder(doc,
        "Examples: 'The Signal · 014 — Maritime tightens, defence loosens'  ·  "
        "'The Signal · 027 — A quiet week in real estate'")

    add_heading(doc, 'Subtitle (Substack subtitle field)', level=2)
    add_placeholder(doc,
        "[One sentence naming the most important shift across the portfolio this week. "
        "No emoji. No hype. Reads like the cover line of an intelligence brief.]")

    add_heading(doc, 'Section structure', level=2)

    add_section_header(doc, 'Hook (1 paragraph, ≤ 60 words)')
    add_placeholder(doc,
        "[State the single most important pattern across the five domains this week. "
        "Plain language. Lead with the noun, not the verb. If you cannot name the pattern in "
        "one sentence, the digest is not ready to publish.]")

    add_section_header(doc, 'Body — five signals')
    add_body(doc,
        "Use the same five-row structure every week. Readers will start scanning for their "
        "domain after issue 3. Do not skip a domain — write 'Quiet week.' if there is nothing "
        "to report. Consistency is the product.")

    add_template_block(doc,
        "## Defence — Aegis\n"
        "[One sentence on what changed in the threat landscape, OSINT volume, or detection "
        "patterns this week. One sentence on what it means for operators.]\n\n"
        "## Maritime — Vessels\n"
        "[One sentence on rates, congestion, or fleet movements. One sentence on the "
        "downstream operational implication.]\n\n"
        "## Real estate — Terra\n"
        "[One sentence on listings, transaction velocity, or distressed-asset signal. "
        "One sentence on what brokers should do about it.]\n\n"
        "## Executive — Command\n"
        "[One sentence on a portfolio-wide KPI shift, board-level concern, or capital event. "
        "One sentence on the C-suite read.]\n\n"
        "## Cross-domain — CORTEX\n"
        "[One observation that only becomes visible when signals from two or more verticals "
        "are correlated. This is the differentiator. Make it count.]")

    add_section_header(doc, 'Closing read (1 paragraph, ≤ 80 words)')
    add_placeholder(doc,
        "[Pull the threads together. What does it mean for a leader making a decision this "
        "week? End on a clean line — not a question, not a cliffhanger.]")

    add_section_header(doc, 'CTA — single line')
    add_template_block(doc,
        "If The Signal is useful, forward it to one operator who would read it. That is how this grows.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings")
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # TEMPLATE 2 — Founder Note
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, 'Template 2 — Founder Note (Essay)', level=1)
    add_body(doc,
        "The essay format. One idea, fully argued. Publishes when you have something to say — "
        "not on a fixed schedule. This is the post readers screenshot and forward.", italic=True)

    add_label_value(doc, 'Cadence', 'Bi-weekly target — never forced')
    add_label_value(doc, 'Target length', '900–1,400 words / 6–9 minute read')
    add_label_value(doc, 'Post tag', 'founder-note')

    add_heading(doc, 'Headline formula', level=2)
    add_body(doc, "Pick one of three patterns. Pattern wins more than wording.")
    add_body(doc, "1. The contrarian claim:  '[Conventional wisdom] is wrong. Here is what we are doing instead.'")
    add_body(doc, "2. The named principle:  '[A short, memorable principle in 2–4 words]'")
    add_body(doc, "3. The earned lesson:  'What [specific decision / shipping moment] taught me about [broader topic]'")
    add_placeholder(doc,
        "Examples: 'Observability before optimisation'  ·  "
        "'Most AI dashboards are theatre. Here is what we built instead.'  ·  "
        "'What shipping CORTEX taught me about audit-grade systems'")

    add_heading(doc, 'Subtitle', level=2)
    add_placeholder(doc,
        "[One sentence stating the thesis as a flat claim — not a question. Should be quotable on its own.]")

    add_heading(doc, 'Section structure', level=2)

    add_section_header(doc, 'Hook (≤ 80 words)')
    add_placeholder(doc,
        "[Open with a concrete moment, decision, or observation. Never start with 'I have been thinking about…'. "
        "Drop the reader inside the scene or the problem. End the hook with the thesis sentence in italics.]")

    add_section_header(doc, 'Body — argument arc (3 sections)')
    add_body(doc,
        "Use three H2 sections. Each one earns the next. Each one is a self-contained argument that could "
        "stand as a tweet on its own.")

    add_template_block(doc,
        "## What everyone agrees on\n"
        "[Establish the conventional view in 100–150 words. Steel-man it. "
        "If you cannot describe the opposing position fairly, you have not earned the right to disagree with it.]\n\n"
        "## Where it breaks\n"
        "[Show the seam. Use one specific example from inside SZL — a shipping decision, a customer "
        "conversation, a system behaviour. Concrete > abstract every time. 200–350 words.]\n\n"
        "## What we do instead\n"
        "[The constructive turn. Name the principle, show how it shows up in the product, and explain "
        "why it produces a better outcome. Reference the specific platform (Aegis, Vessels, Terra, "
        "Command, CORTEX, or Alloy) where the principle lives. 250–400 words.]")

    add_section_header(doc, 'Close (≤ 100 words)')
    add_placeholder(doc,
        "[State the principle one more time, in its sharpest form. No hedging. No 'of course this is "
        "just my opinion.' If you do not believe it firmly enough to defend, do not publish it.]")

    add_section_header(doc, 'CTA — single line')
    add_template_block(doc,
        "If you disagree, reply to this email. I read every one.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings")
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # TEMPLATE 3 — Product / Build Log
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, 'Template 3 — Build Log (Product Update)', level=1)
    add_body(doc,
        "What shipped, what changed, what is next. The proof-over-pitch post. "
        "Publishes whenever a meaningful capability lands in any of the five platforms.", italic=True)

    add_label_value(doc, 'Cadence', 'On every meaningful release — typically 2–3× per month')
    add_label_value(doc, 'Target length', '500–800 words / 3–5 minute read')
    add_label_value(doc, 'Post tag', 'build-log')

    add_heading(doc, 'Headline formula', level=2)
    add_body(doc, "[Platform name] — [verb in past tense] [the thing]")
    add_placeholder(doc,
        "Examples: 'Aegis — shipped one-click incident response across the SOC queue'  ·  "
        "'Terra — replaced the inquiry inbox with routed broker queues'  ·  "
        "'CORTEX — correlated signals are now scored end-to-end'")

    add_heading(doc, 'Subtitle', level=2)
    add_placeholder(doc,
        "[One sentence naming who this matters to and why now. e.g. "
        "'For SOC leads tired of swivel-chairing between five tools to close one ticket.']")

    add_heading(doc, 'Section structure', level=2)

    add_section_header(doc, 'Hook — the why (≤ 70 words)')
    add_placeholder(doc,
        "[Name the operational pain this release solves. One sentence on the user. One sentence on what "
        "they were doing before. No product marketing.]")

    add_section_header(doc, 'Body — what shipped')
    add_template_block(doc,
        "## What is new\n"
        "[Plain-language description of the capability. 2–4 short paragraphs. "
        "Use product names in italics on first reference. If a screenshot or short video clip "
        "would replace 100 words, embed it instead of writing the words.]\n\n"
        "## How it works\n"
        "[The architecture in one paragraph. Name the data sources, the action layer (Alloy), "
        "and the surface where it appears. Operators want to know how it routes, not how it markets.]\n\n"
        "## What it changes for operators\n"
        "[The before/after, in concrete terms. Number of clicks. Time to resolution. Steps eliminated. "
        "If you cannot quantify the change, you may be shipping a feature, not a capability.]")

    add_section_header(doc, 'What is next (≤ 80 words)')
    add_placeholder(doc,
        "[The one thing landing next on this platform. Specific. Date-bounded if possible. "
        "Builds the publishing flywheel — readers come back to see if you shipped what you said you would.]")

    add_section_header(doc, 'CTA — single line')
    add_template_block(doc,
        "If you operate in this space and want a working session on it, reply to this email.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings")
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # TEMPLATE 4 — Field Report (Decision Teardown)
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, 'Template 4 — Field Report (Decision Teardown / Post-Mortem)', level=1)
    add_body(doc,
        "An honest account of one decision: what happened, what was decided, what it cost, "
        "and what we learned. The post that builds long-term trust. "
        "Publishes monthly at minimum. Earns the right to charge for paid posts later.", italic=True)

    add_label_value(doc, 'Cadence', 'Monthly — last Friday of the month')
    add_label_value(doc, 'Target length', '700–1,100 words / 5–7 minute read')
    add_label_value(doc, 'Post tag', 'field-report')

    add_heading(doc, 'Headline formula', level=2)
    add_body(doc, "Field report — [the decision, named flatly]")
    add_placeholder(doc,
        "Examples: 'Field report — why we cut the second OSINT vendor'  ·  "
        "'Field report — how we lost two weeks on the wrong abstraction'  ·  "
        "'Field report — what the first design partner taught us about pricing'")

    add_heading(doc, 'Subtitle', level=2)
    add_placeholder(doc,
        "[One sentence stating the outcome plainly. 'We were wrong about X.' or 'Y cost us Z weeks.' "
        "Honesty is the format.]")

    add_heading(doc, 'Section structure', level=2)

    add_section_header(doc, 'Hook — the moment (≤ 80 words)')
    add_placeholder(doc,
        "[Open in the room where the decision was made. Who was there. What was on the table. "
        "What you knew and did not know at that moment. No retrospective varnish.]")

    add_section_header(doc, 'Body — the four-part teardown')
    add_template_block(doc,
        "## What we believed going in\n"
        "[The thesis at decision time. State it the way you would have written it the morning of. "
        "100–200 words.]\n\n"
        "## What actually happened\n"
        "[The events, in order. Dates if you have them. Numbers if you have them. Be specific enough "
        "that a reader could audit the account. 200–350 words.]\n\n"
        "## What it cost\n"
        "[Time, capital, opportunity, trust. Quantify what can be quantified. Acknowledge what cannot. "
        "Do not soften the number. 100–200 words.]\n\n"
        "## What we changed because of it\n"
        "[The concrete operational change. A new check in the workflow, a default flipped, a hire we "
        "now make earlier, a question we now ask first. The change is the proof the lesson landed. "
        "150–250 words.]")

    add_section_header(doc, 'Close (≤ 60 words)')
    add_placeholder(doc,
        "[One sentence on what you would tell a founder facing the same decision next week. "
        "Not 'be careful.' Specific.]")

    add_section_header(doc, 'CTA — single line')
    add_template_block(doc,
        "If you have made the same call and learned something different, reply. I want to hear it.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings")
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # TEMPLATE 5 — Curated Reading List
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, 'Template 5 — On the Desk (Curated Reading & Watchlist)', level=1)
    add_body(doc,
        "What is on the operator's desk this month: the reports, dashboards, papers, and tools "
        "actually informing decisions across the portfolio. Light to write, heavy to read. "
        "Drives subscriber growth — readers forward this format more than any other.", italic=True)

    add_label_value(doc, 'Cadence', 'Monthly — first Friday of the month')
    add_label_value(doc, 'Target length', '500–750 words / 3–4 minute read')
    add_label_value(doc, 'Post tag', 'on-the-desk')

    add_heading(doc, 'Headline formula', level=2)
    add_body(doc, "On the desk — [Month] [Year]")
    add_placeholder(doc,
        "Examples: 'On the desk — April 2026'  ·  'On the desk — Q2 opening'")

    add_heading(doc, 'Subtitle', level=2)
    add_placeholder(doc,
        "[One sentence naming the through-line across this month's items. e.g. "
        "'Five reads on why audit-grade is becoming table stakes for enterprise AI.']")

    add_heading(doc, 'Section structure', level=2)

    add_section_header(doc, 'Hook (≤ 50 words)')
    add_placeholder(doc,
        "[State the through-line. Why these items, this month. One sentence. No 'I have been reading a lot lately'.]")

    add_section_header(doc, 'Body — five items, same shape every time')
    add_body(doc,
        "Five items. One link, one read, one reason. Same structure every month so readers learn the rhythm.")

    add_template_block(doc,
        "## 1. [Title of the report / paper / dashboard / tool]\n"
        "[Source · author · date.] [Link.]\n"
        "[Two sentences: what it is, and the one insight worth keeping. End with: "
        "'Why it is on the desk: [one specific reason tied to a portfolio decision].']\n\n"
        "## 2. [Title]\n"
        "[Source · author · date.] [Link.]\n"
        "[Two sentences + 'Why it is on the desk' line.]\n\n"
        "## 3. [Title]\n"
        "[Source · author · date.] [Link.]\n"
        "[Two sentences + 'Why it is on the desk' line.]\n\n"
        "## 4. [Title]\n"
        "[Source · author · date.] [Link.]\n"
        "[Two sentences + 'Why it is on the desk' line.]\n\n"
        "## 5. [Title]\n"
        "[Source · author · date.] [Link.]\n"
        "[Two sentences + 'Why it is on the desk' line.]")

    add_section_header(doc, 'Close — one line')
    add_placeholder(doc,
        "[Name the one item on this list a reader should open first if they only open one.]")

    add_section_header(doc, 'CTA — single line')
    add_template_block(doc,
        "Reply with what is on your desk this month. The best ones land in the next issue.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings")
    add_divider(doc)

    # ── Publishing rhythm ─────────────────────────────────────────────────────
    add_heading(doc, 'Publishing Rhythm — A Default Month', level=1)
    add_body(doc,
        "Use the cadence below as a starting point. The five templates compose into "
        "8–10 posts per month without writer's block — each one with a fixed shape so the "
        "writing time goes into substance, not structure.")

    add_label_value(doc, 'Every Monday', 'Template 1 — The Signal (4× per month)')
    add_label_value(doc, '1st Friday', 'Template 5 — On the Desk')
    add_label_value(doc, 'Mid-month', 'Template 2 — Founder Note (when the idea is ready)')
    add_label_value(doc, '2–3× per month', 'Template 3 — Build Log (whenever a release lands)')
    add_label_value(doc, 'Last Friday', 'Template 4 — Field Report')

    add_note(doc,
        "Never miss a Monday. The Signal is the contract with the reader. If a week is genuinely "
        "quiet, write the shortest possible Signal — never skip it.")
    add_divider(doc)

    # ── How to use ────────────────────────────────────────────────────────────
    add_heading(doc, 'How to Use This Document', level=1)
    add_body(doc, "1. Open this file in Word or Google Docs alongside Substack's post editor.")
    add_body(doc, "2. Pick the template that matches what you have to say this week.")
    add_body(doc, "3. Copy the headline formula into Substack's title field and replace the bracketed parts.")
    add_body(doc, "4. Copy the subtitle line into Substack's subtitle field.")
    add_body(doc, "5. Copy each section block into the post body in order. Replace every bracketed "
                  "placeholder with your own content — never publish with brackets remaining.")
    add_body(doc, "6. Run the Voice & Brand Notes checklist at the top of this document before hitting publish.")
    add_body(doc, "7. Apply the standard post settings (Free · Newsletter · Comments on).")
    add_body(doc, "8. Save the published URL into your editorial log so future Build Logs can link back.")

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
