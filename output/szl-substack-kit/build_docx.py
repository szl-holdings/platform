#!/usr/bin/env python3
"""Generate SZL-Substack-Profile.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "/home/runner/workspace/output/szl-substack-kit/SZL-Substack-Profile.docx"

# Brand colors (converted from HSL to RGB approximations)
GOLD = RGBColor(0xD4, 0xA0, 0x54)       # hsl(38,52%,58%) — SZL accent gold
NAVY_DARK = RGBColor(0x07, 0x0A, 0x10)  # background (not used in docx, dark mode)
PLATINUM = RGBColor(0xB8, 0xC2, 0xCC)   # hsl(210,8%,78%)
BLUE = RGBColor(0x3D, 0x6F, 0xD9)       # szl-primary blue
DARK_GREY = RGBColor(0x22, 0x2B, 0x38)  # deep slate for body text on white
MUTED = RGBColor(0x64, 0x74, 0x88)      # secondary text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_heading_color(run, color: RGBColor):
    run.font.color.rgb = color


def add_heading(doc, text, level=1, color=GOLD):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(20)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = BLUE
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = DARK_GREY
    return h


def add_label_value(doc, label, value, note=None):
    """Add a field label + value pair."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = DARK_GREY
    run_label.font.size = Pt(11)
    run_value = p.add_run(value)
    run_value.font.color.rgb = BLACK
    run_value.font.size = Pt(11)
    if note:
        run_note = p.add_run(f"  ({note})")
        run_note.font.color.rgb = MUTED
        run_note.font.size = Pt(10)
        run_note.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GREY
        if italic:
            run.font.italic = True
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"Note: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run.font.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D4A054')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    return p


def build_doc():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Cover title block ──────────────────────────────────────────────────────
    title = doc.add_heading('SZL Holdings', 0)
    for run in title.runs:
        run.font.color.rgb = GOLD
        run.font.size = Pt(28)
        run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Substack Profile Kit — Stephen Lutar')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = MUTED
        run.font.size = Pt(12)
        run.font.italic = True
    doc.add_paragraph()

    add_body(doc,
        "This document contains every piece of copy you need to fully populate your Substack "
        "profile at szlholdings.substack.com. Copy each field directly into Edit Profile — "
        "no editing required unless you want to personalise further.",
        italic=True)
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 1: PROFILE IDENTITY
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '1. Profile Identity', level=1)

    add_label_value(doc, 'Name', 'Stephen Lutar')
    add_label_value(doc, 'Handle / Subdomain', 'szlholdings',
        note='Profile URL will be szlholdings.substack.com')
    add_label_value(doc, 'Publication Name', 'SZL Holdings')
    add_label_value(doc, 'Tagline', 'Structured ventures. Clear direction.')
    add_label_value(doc, 'Location', 'New York, NY')
    add_label_value(doc, 'Website', 'https://szlholdings.com')
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 2: BIO VARIANTS
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '2. Bio Variants', level=1)
    add_body(doc,
        "Substack surfaces the bio in multiple contexts (profile page, about section, email footer). "
        "Choose the variant that fits the context, or use them in order as the platform allows.", italic=True)

    add_heading(doc, '2a. Short Bio (≤ 160 characters — for profile tagline / search snippet)', level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Founder & CEO at SZL Holdings. Aegis · Vessels · Terra · Command · CORTEX — "
        "intelligence systems for high-stakes operations."
    ).font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    add_note(doc, "150 characters. Use exactly as written in the 'About' tagline field.")

    add_heading(doc, '2b. Medium Bio (≤ 300 characters — for newsletter description)', level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Stephen Lutar is the Founder & CEO of SZL Holdings — a portfolio of intelligence platforms "
        "spanning defense (Aegis), maritime (Vessels), real estate (Terra), executive command (Command), "
        "and cross-domain intelligence (CORTEX). Builder-operator. Proof over pitch. Based in New York."
    ).font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    add_note(doc, "298 characters. Use in the 'Publication description' field on Edit Publication.")

    add_heading(doc, '2c. Long Bio (≤ 500 words — for About page / full profile)', level=2)
    long_bio = (
        "Stephen Lutar is the Founder & CEO of SZL Holdings, a vertically structured portfolio of "
        "intelligence and operations platforms built for organisations where decisions carry real weight.\n\n"
        "The portfolio spans five primary platforms. Aegis delivers unified threat intelligence, SOC command, "
        "and incident response orchestration for defence and national-security environments. Vessels provides "
        "maritime intelligence and fleet operations command for shipping operators and port authorities. "
        "Terra is the broker-grade intelligence and CRM layer for commercial real estate — active listings, "
        "market signals, and inquiry routing in one command surface. Command is the unified executive "
        "command portal: portfolio intelligence, KPI dashboards, and strategic briefings for the C-suite "
        "and board of directors. And CORTEX is the cross-domain intelligence hub that correlates signals "
        "across every vertical, giving leadership a single, audit-grade picture of the entire portfolio.\n\n"
        "Beneath every platform sits Alloy, the shared execution fabric that handles ingestion, "
        "normalisation, and workflow automation — ensuring that intelligence doesn't just surface; "
        "it routes, escalates, and closes the loop with a verifiable record.\n\n"
        "Stephen operates as a builder-operator: he builds the product, runs the architecture, and "
        "works directly with design partners. There is no separation between vision and execution. "
        "Every claim is backed by running code and live systems. The default is always to show a "
        "working product rather than describe a future one.\n\n"
        "Operating philosophy: observability before optimisation. Proof over pitch. "
        "Trust-first architecture built from day one — not retrofitted when enterprise buyers demand it.\n\n"
        "SZL Holdings is based in New York, NY."
    )
    p = doc.add_paragraph(long_bio)
    for run in p.runs:
        run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    add_note(doc, "Use this text in the full 'About' page body on your publication settings.")
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 3: SOCIAL LINKS
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '3. Social Links', level=1)
    add_body(doc,
        "Enter the following values exactly as shown in the Social links section of Edit Profile. "
        "Fields marked 'Leave blank' should remain empty.", italic=True)

    add_label_value(doc, 'Website', 'https://szlholdings.com')
    add_label_value(doc, 'X (Twitter)', 'https://x.com/szlholdings')
    add_label_value(doc, 'LinkedIn', 'https://linkedin.com/in/stephenlutar')
    add_label_value(doc, 'Medium', 'https://medium.com/@stephen_38454')
    add_label_value(doc, 'Instagram', '', note='Leave blank')
    add_label_value(doc, 'Facebook', '', note='Leave blank')
    add_label_value(doc, 'YouTube', '', note='Leave blank')
    add_label_value(doc, 'TikTok', '', note='Leave blank')
    add_label_value(doc, 'GitHub', '', note='Leave blank — not public-facing at this stage')
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 4: SETTINGS
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '4. Recommended Settings', level=1)

    add_heading(doc, 'Writes / Publication type', level=2)
    add_label_value(doc, 'Publication type', 'Newsletter + Publication',
        note='Set to allow both email posts and web archive')
    add_label_value(doc, 'Default access level', 'Free',
        note='Maximises reach during the audience-building phase; gate premium content when ready')
    add_label_value(doc, 'Comments', 'On — all subscribers',
        note='Encourages engagement and signals activity to new visitors')

    add_heading(doc, 'Subscription Privacy', level=2)
    add_label_value(doc, 'Hide subscriber count', 'No — show publicly',
        note='Once count grows, social proof accelerates growth; start transparent')
    add_label_value(doc, 'Hide following list', 'Yes',
        note='Protects your reading habits and signal sourcing')

    add_heading(doc, 'Activity Privacy', level=2)
    add_label_value(doc, 'Show activity on profile', 'Yes — likes and restacks visible',
        note='Demonstrates active presence on the platform')
    add_label_value(doc, 'Show what you are subscribed to', 'No — keep private',
        note='Competitive intelligence hygiene')

    add_heading(doc, 'Subscriber Badge', level=2)
    add_label_value(doc, 'Enable founding member badge', 'Yes',
        note='Creates early-adopter identity for your first subscribers')
    add_label_value(doc, 'Founding member price', 'Set when ready to monetise',
        note='Leave at default until the first paid offering is ready')

    add_heading(doc, 'Theme', level=2)
    add_label_value(doc, 'Recommended theme', 'Dark / System (respects reader preference)',
        note='Aligns with SZL brand aesthetic; most professional readers use dark mode')
    add_label_value(doc, 'Accent colour', '#D4A054',
        note='SZL Holdings gold — enter as hex in the Custom colour field')
    add_label_value(doc, 'Font pairing', 'Sans-serif display + body',
        note='Select the most neutral sans-serif option available; Space Grotesk is not available on Substack but "Ideal Sans" or the default Sans works well')
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 5: COVER IMAGE GUIDANCE
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '5. Cover Image Guidance', level=1)
    add_body(doc,
        "Two cover image files are included in this kit. Upload the 1600×533 version first — "
        "Substack will downsample it for standard displays. If Substack requests a smaller file, "
        "use the 1200×400 version.", italic=True)

    add_label_value(doc, 'Standard cover', 'cover-image-1200x400.png  (1200 × 400 px, 3:1 ratio)')
    add_label_value(doc, 'Retina / high-DPI cover', 'cover-image-1600x533.png  (1600 × 533 px, 3:1 ratio)')
    add_label_value(doc, 'Background', 'Deep navy to black gradient — aligns with SZL Holdings dark brand')
    add_label_value(doc, 'Wordmark', 'SZL Holdings in gold (#D4A054) with platinum tagline')
    add_label_value(doc, 'Upload location', 'Substack Dashboard → Edit Publication → Cover image')
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 6: WELCOME / INTRO POST
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '6. Welcome Post (Pinned Intro)', level=1)
    add_body(doc,
        "Copy the text below as your first and pinned post. "
        "Publish it as a free post visible to all visitors.", italic=True)
    add_label_value(doc, 'Suggested post title', 'What SZL Holdings Is Building — and Why')
    add_label_value(doc, 'Post type', 'Free — visible to everyone')
    add_label_value(doc, 'Pin to top', 'Yes')

    doc.add_paragraph()
    add_heading(doc, 'Post body (copy verbatim):', level=3)
    welcome_post = (
        "Most companies talk about intelligence. We build the systems that make it operational.\n\n"
        "SZL Holdings is a vertically integrated portfolio of intelligence platforms — each one built "
        "for a domain where decisions have real consequences: defence, maritime, real estate, "
        "executive command, and cross-domain intelligence.\n\n"
        "Here is what we have shipped:\n\n"
        "Aegis — threat intelligence fusion and SOC command for defence and enterprise security. "
        "Ingests 900+ OSINT and ISAC feeds, correlates them against your internal telemetry, and "
        "surfaces ranked, explainable threats with one-click response.\n\n"
        "Vessels — maritime intelligence and fleet operations command. Route optimisation, port authority "
        "interfaces, cargo intelligence, and voyage P&L in one command surface built for how shipping "
        "operators actually work.\n\n"
        "Terra — real estate broker command. Active listings, market signals, pre-foreclosure data, "
        "and inquiry routing — not a consumer portal, a precision tool for commercial brokers.\n\n"
        "Command — the unified executive command portal. Portfolio intelligence, KPI dashboards, and "
        "strategic briefings for the C-suite and board of directors, pulled from every operating company "
        "in the portfolio in real time.\n\n"
        "CORTEX — the cross-domain intelligence hub. Every signal from every vertical converges here "
        "for executive leadership. Audit-grade, explainable, and wired to the action layer.\n\n"
        "This Substack is where I publish what I am learning building these systems — "
        "architecture decisions, operational patterns, and the occasional honest post-mortem.\n\n"
        "No pitch. No noise. Just the work.\n\n"
        "— Stephen Lutar, Founder & CEO, SZL Holdings"
    )
    p = doc.add_paragraph(welcome_post)
    for run in p.runs:
        run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    add_divider(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 7: QUICK REFERENCE
    # ──────────────────────────────────────────────────────────────────────────
    add_heading(doc, '7. Quick Reference Card', level=1)
    add_body(doc, "Everything in one glance:", italic=True)

    fields = [
        ("Publication name", "SZL Holdings"),
        ("Handle", "szlholdings"),
        ("Tagline", "Structured ventures. Clear direction."),
        ("Author name", "Stephen Lutar"),
        ("Author title", "Founder & CEO, SZL Holdings"),
        ("Website", "https://szlholdings.com"),
        ("X / Twitter", "https://x.com/szlholdings"),
        ("LinkedIn", "https://linkedin.com/in/stephenlutar"),
        ("Medium", "https://medium.com/@stephen_38454"),
        ("Accent colour", "#D4A054"),
        ("Cover image", "cover-image-1600x533.png (upload this one)"),
    ]
    for label, value in fields:
        add_label_value(doc, label, value)

    doc.add_paragraph()
    add_body(doc,
        "All files included in this kit: SZL-Substack-Profile.docx, cover-image-1200x400.png, "
        "cover-image-1600x533.png, profile-mockup-desktop.png, profile-mockup-mobile.png, README.txt",
        italic=True)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
