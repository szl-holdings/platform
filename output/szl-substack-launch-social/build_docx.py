#!/usr/bin/env python3
"""Generate SZL-Substack-Launch-Social-Kit.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "/home/runner/workspace/output/szl-substack-launch-social/SZL-Substack-Launch-Social-Kit.docx"

GOLD = RGBColor(0xD4, 0xA0, 0x54)
BLUE = RGBColor(0x3D, 0x6F, 0xD9)
DARK_GREY = RGBColor(0x22, 0x2B, 0x38)
MUTED = RGBColor(0x64, 0x74, 0x88)
BLACK = RGBColor(0x00, 0x00, 0x00)


def add_heading(doc, text, level=1, color=GOLD):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(20)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = BLUE
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = DARK_GREY
    return h


def add_label_value(doc, label, value, note=None):
    p = doc.add_paragraph()
    rl = p.add_run(f"{label}: ")
    rl.bold = True
    rl.font.color.rgb = DARK_GREY
    rl.font.size = Pt(11)
    rv = p.add_run(value)
    rv.font.color.rgb = BLACK
    rv.font.size = Pt(11)
    if note:
        rn = p.add_run(f"  ({note})")
        rn.font.color.rgb = MUTED
        rn.font.size = Pt(10)
        rn.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GREY
        if italic:
            run.font.italic = True
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"Note: {text}")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    r.font.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_copy_block(doc, text, char_count_label=None):
    """Indented monospace-feeling copy block ready to copy/paste."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    if char_count_label:
        add_note(doc, char_count_label)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D4A054')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    return p


def cc(s):
    """Char count helper label."""
    return f"{len(s)} characters."


def build_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Cover
    title = doc.add_heading('SZL Holdings', 0)
    for run in title.runs:
        run.font.color.rgb = GOLD
        run.font.size = Pt(28)
        run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Substack Launch — Social Announcement Kit')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = MUTED
        run.font.size = Pt(12)
        run.font.italic = True
    doc.add_paragraph()

    add_body(doc,
        "Ready-to-post copy for announcing szlholdings.substack.com across X (Twitter) and "
        "LinkedIn, plus a personal email for existing contacts and design partners. Every "
        "block is written in the SZL voice — proof over pitch, no noise — and is sized to "
        "the platform it belongs on. Copy verbatim or personalise lightly.",
        italic=True)
    add_divider(doc)

    # ─────────────────────────────────────────────────────────────
    # SECTION 1: X / TWITTER THREAD STARTERS
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, '1. X (Twitter) Thread Starters', level=1)
    add_body(doc,
        "Three opening tweets for an announcement thread. Each is under 280 characters and "
        "stands on its own as a single post if you do not want to thread. Pick one. The "
        "follow-up replies underneath each starter are optional — use them if you want to "
        "extend it into a 3-tweet thread.",
        italic=True)

    # Variant A
    add_heading(doc, 'Variant A — The Direct Announce', level=2)
    a1 = ("New: szlholdings.substack.com is live.\n\n"
          "Architecture decisions, operational patterns, and the occasional honest "
          "post-mortem from building Aegis, Vessels, Terra, Command, and CORTEX.\n\n"
          "No pitch. No noise. Just the work.\n\n"
          "Subscribe → szlholdings.substack.com")
    add_copy_block(doc, a1, cc(a1))

    add_heading(doc, 'Optional follow-up replies (Variant A):', level=3)
    a2 = ("What you will get:\n\n"
          "— The Signal: weekly intelligence digest\n"
          "— Founder Note: essays on building intelligence systems\n"
          "— Build Log: what shipped and why\n"
          "— Field Report: monthly decision teardowns\n\n"
          "Free, for now.")
    add_copy_block(doc, a2, cc(a2))
    a3 = ("Why I am writing in public:\n\n"
          "Most companies talk about intelligence. Few build the systems that make it "
          "operational. I want to show the work — the decisions, the trade-offs, and the "
          "post-mortems — for anyone building in the same direction.")
    add_copy_block(doc, a3, cc(a3))

    # Variant B
    add_heading(doc, 'Variant B — The Hook', level=2)
    b1 = ("Most companies talk about intelligence.\n\n"
          "We build the systems that make it operational.\n\n"
          "Starting today I am publishing how — architecture, operations, and post-mortems "
          "from inside SZL Holdings.\n\n"
          "→ szlholdings.substack.com")
    add_copy_block(doc, b1, cc(b1))

    add_heading(doc, 'Optional follow-up replies (Variant B):', level=3)
    b2 = ("Five platforms inside the portfolio:\n\n"
          "Aegis — defence threat intelligence\n"
          "Vessels — maritime fleet command\n"
          "Terra — broker-grade real estate intelligence\n"
          "Command — executive command portal\n"
          "CORTEX — cross-domain intelligence hub")
    add_copy_block(doc, b2, cc(b2))
    b3 = ("Cadence:\n\n"
          "Weekly signal digest. Bi-weekly founder essay. Build log on every meaningful "
          "release.\n\n"
          "If you build, operate, or invest in high-stakes systems, this is for you.\n\n"
          "Subscribe → szlholdings.substack.com")
    add_copy_block(doc, b3, cc(b3))

    # Variant C
    add_heading(doc, 'Variant C — The Builder-Operator Note', level=2)
    c1 = ("I write code in the morning, run the architecture in the afternoon, and meet with "
          "design partners at night.\n\n"
          "Now I am writing about it: szlholdings.substack.com\n\n"
          "Architecture decisions, operational patterns, honest post-mortems. Free.")
    add_copy_block(doc, c1, cc(c1))

    add_heading(doc, 'Optional follow-up replies (Variant C):', level=3)
    c2 = ("Operating philosophy in three lines:\n\n"
          "Observability before optimisation.\n"
          "Proof over pitch.\n"
          "Trust-first architecture from day one — never retrofitted.\n\n"
          "If that resonates, you will like what is coming.")
    add_copy_block(doc, c2, cc(c2))
    c3 = ("First posts dropping over the next two weeks:\n\n"
          "— What SZL Holdings Is Building (and Why)\n"
          "— The Signal #001\n"
          "— Build Log: shipping CORTEX\n\n"
          "Subscribe so you do not miss them: szlholdings.substack.com")
    add_copy_block(doc, c3, cc(c3))

    add_divider(doc)

    # ─────────────────────────────────────────────────────────────
    # SECTION 2: LINKEDIN POSTS
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, '2. LinkedIn Posts', level=1)
    add_body(doc,
        "Two variants. Variant A is short and sharp — best for a busy feed. Variant B is the "
        "longer essay-style announcement, better for organic reach and reposts. Pick one or "
        "post both, two weeks apart. LinkedIn permits up to ~3,000 characters per post; "
        "both are well within that limit.",
        italic=True)

    add_heading(doc, 'Variant A — Short Post', level=2)
    la = ("Today I am launching a Substack: szlholdings.substack.com\n\n"
          "It is where I will publish what I am learning building SZL Holdings — a "
          "vertically integrated portfolio of intelligence platforms spanning defence "
          "(Aegis), maritime (Vessels), real estate (Terra), executive command (Command), "
          "and cross-domain intelligence (CORTEX).\n\n"
          "Expect:\n\n"
          "— Architecture decisions and operational patterns\n"
          "— Weekly intelligence signal digests\n"
          "— Build logs on every meaningful release\n"
          "— Honest post-mortems when something does not work\n\n"
          "No pitch. No noise. Just the work.\n\n"
          "If you build, operate, or invest in systems where decisions carry real "
          "consequences, this should be in your inbox.\n\n"
          "Subscribe → szlholdings.substack.com")
    add_copy_block(doc, la, cc(la))

    add_heading(doc, 'Variant B — Longer Essay-Style Post', level=2)
    lb = (
        "Most companies talk about intelligence. Far fewer build the systems that make it "
        "operational.\n\n"
        "For the last several years I have been doing the second thing — quietly, with "
        "design partners, inside SZL Holdings. Today I am opening that work to a wider "
        "audience.\n\n"
        "szlholdings.substack.com is now live.\n\n"
        "A short tour of what we have built:\n\n"
        "Aegis — unified threat intelligence and SOC command for defence and enterprise "
        "security. 900+ OSINT and ISAC feeds correlated against internal telemetry, with "
        "ranked, explainable threats and one-click response.\n\n"
        "Vessels — maritime intelligence and fleet operations command. Route optimisation, "
        "port authority interfaces, cargo intelligence, and voyage P&L in one command "
        "surface.\n\n"
        "Terra — broker-grade real estate intelligence. Active listings, market signals, "
        "pre-foreclosure data, and inquiry routing for commercial brokers.\n\n"
        "Command — the unified executive command portal for the C-suite and board.\n\n"
        "CORTEX — the cross-domain intelligence hub that converges signals from every "
        "vertical into a single, audit-grade picture.\n\n"
        "Beneath all of it sits Alloy — the shared execution fabric that ingests, "
        "normalises, and routes work so intelligence does not just surface; it closes the "
        "loop with a verifiable record.\n\n"
        "On the Substack I will publish what is hard to fit into a product page:\n\n"
        "— Why we made the architecture decisions we made\n"
        "— What the operating patterns look like in practice\n"
        "— What shipped, what broke, and what we changed\n"
        "— Weekly intelligence signal digests across the verticals we cover\n\n"
        "Operating philosophy, three lines:\n"
        "Observability before optimisation.\n"
        "Proof over pitch.\n"
        "Trust-first architecture, built from day one — never retrofitted when enterprise "
        "buyers ask.\n\n"
        "If you build, operate, or invest in high-stakes systems, this should be useful.\n\n"
        "Subscribe → szlholdings.substack.com\n\n"
        "And if there is a topic you want me to cover early, leave a comment — I read "
        "every one."
    )
    add_copy_block(doc, lb, cc(lb))

    add_divider(doc)

    # ─────────────────────────────────────────────────────────────
    # SECTION 3: PERSONAL EMAIL
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, '3. Personal Email — Existing Contacts & Design Partners', level=1)
    add_body(doc,
        "A short, direct email for sending one-to-one (or in small BCC batches) to existing "
        "contacts, design partners, prior colleagues, and warm investors. Personalise the "
        "first line where time permits — it doubles open and click rates. Send from your "
        "personal address, not a marketing domain.",
        italic=True)

    add_label_value(doc, 'Subject line (Option 1)', 'Quietly launching something — would value your eye on it')
    add_label_value(doc, 'Subject line (Option 2)', 'New: szlholdings.substack.com')
    add_label_value(doc, 'Subject line (Option 3)', 'A short note — and a link')
    add_label_value(doc, 'From',  'Stephen Lutar <your personal address>')
    add_label_value(doc, 'To',    'One recipient at a time, or small BCC batches (≤ 30)')

    doc.add_paragraph()
    add_heading(doc, 'Email body (copy verbatim, edit the first line):', level=3)
    email = (
        "Hi {first name},\n\n"
        "{One personal line — last time we spoke / something they recently shipped / "
        "a shared connection. Two sentences max.}\n\n"
        "Quick note: I have started writing publicly about what we are building inside "
        "SZL Holdings. The Substack is live at szlholdings.substack.com.\n\n"
        "It will cover the architecture decisions, operating patterns, and post-mortems "
        "behind Aegis (defence), Vessels (maritime), Terra (real estate), Command "
        "(executive portal), and CORTEX (cross-domain intelligence) — plus a weekly "
        "signal digest across the domains we cover.\n\n"
        "No pitch, no fundraising ask. I would value your eye on it, and any feedback on "
        "the topics that would be most useful to you.\n\n"
        "If it is for you, the subscribe button is in the top right. If it is not, no "
        "harm done — and I would still love to hear what you are working on these days.\n\n"
        "Best,\n"
        "Stephen\n\n"
        "—\n"
        "Stephen Lutar\n"
        "Founder & CEO, SZL Holdings\n"
        "szlholdings.com  ·  szlholdings.substack.com"
    )
    add_copy_block(doc, email)
    add_note(doc,
        "Replace {first name} and the bracketed personal line. Everything else can stay as "
        "written. Send in batches of ≤ 30 BCCs to avoid spam filtering, or one-to-one for "
        "your top 20 relationships.")

    add_divider(doc)

    # ─────────────────────────────────────────────────────────────
    # SECTION 4: HASHTAG & TIMING RECOMMENDATIONS
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, '4. Hashtag & Timing Recommendations', level=1)

    add_heading(doc, 'Hashtags — X (Twitter)', level=2)
    add_body(doc,
        "On X, hashtags are largely cosmetic — the algorithm weights links and engagement, "
        "not tags. Use one or two only, at the end of the tweet, and only when they fit "
        "the audience naturally. Skip them entirely on the opening tweet of a thread.",
        italic=True)
    add_label_value(doc, 'Primary tags',   '#Substack  #BuildingInPublic')
    add_label_value(doc, 'Vertical tags (use 1, only if relevant to the post)',
        '#Defense  #Maritime  #CRE  #ThreatIntel  #IntelligenceSystems')
    add_label_value(doc, 'Avoid',
        '#Newsletter, #Marketing, #ContentCreator',
        note='Wrong audience signal for this brand')

    add_heading(doc, 'Hashtags — LinkedIn', level=2)
    add_body(doc,
        "LinkedIn weights hashtags more than X — use three to five, placed at the very end "
        "of the post (not inline). Mix one broad, one medium, and one niche tag.",
        italic=True)
    add_label_value(doc, 'Recommended set (Variant A)',
        '#Substack  #BuildingInPublic  #IntelligenceSystems  #EnterpriseSoftware')
    add_label_value(doc, 'Recommended set (Variant B)',
        '#FounderJournal  #ThreatIntelligence  #MaritimeTech  #CommercialRealEstate  #BuildingInPublic')
    add_label_value(doc, 'Avoid',
        '#Hustle, #Grindset, #Entrepreneur, #Motivation',
        note='Off-brand for SZL voice')

    add_heading(doc, 'Posting Times', level=2)
    add_body(doc,
        "All times are New York time (ET). These windows are based on B2B engagement "
        "patterns for executives and operators in defence, finance, and real estate. "
        "Decision-makers are at their desks early; sleep-deprived founders scroll late.",
        italic=True)

    add_label_value(doc, 'X — best windows',
        'Tue / Wed / Thu, 7:30–9:30 AM ET  ·  12:00–1:00 PM ET  ·  9:00–10:30 PM ET')
    add_label_value(doc, 'LinkedIn — best windows',
        'Tue / Wed / Thu, 7:30–9:00 AM ET  ·  11:30 AM–12:30 PM ET',
        note='Avoid Friday afternoons and weekends — engagement falls off a cliff')
    add_label_value(doc, 'Email — best send time',
        'Tuesday or Wednesday, 6:45–7:30 AM ET',
        note='Lands at the top of the inbox before the workday starts')

    add_heading(doc, 'Suggested Launch Sequence (one week)', level=2)
    add_body(doc, "Sequencing the channels avoids cannibalising your own announcement.", italic=True)
    add_label_value(doc, 'Day 1 — Tuesday, 7:30 AM ET',
        'Post LinkedIn Variant B (the long essay). Highest reach window of the week.')
    add_label_value(doc, 'Day 1 — Tuesday, 9:00 AM ET',
        'Post X Variant A (Direct Announce) with the optional follow-up replies. Pin to your profile.')
    add_label_value(doc, 'Day 2 — Wednesday, 7:00 AM ET',
        'Send the personal email to your top 30 relationships, one-to-one or small BCC.')
    add_label_value(doc, 'Day 4 — Friday, 8:30 AM ET',
        'Post X Variant C (Builder-Operator Note) — different angle, no thread.')
    add_label_value(doc, 'Day 7 — following Tuesday, 12:00 PM ET',
        'Post LinkedIn Variant A (Short) and X Variant B (The Hook) on the same day for the second wave.')
    add_label_value(doc, 'Ongoing',
        'Restack each new Substack post to X and LinkedIn within an hour of publishing. Keep the cadence visible.')

    add_divider(doc)

    # ─────────────────────────────────────────────────────────────
    # SECTION 5: QUICK REFERENCE
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, '5. Quick Reference', level=1)
    fields = [
        ("Substack URL",       "https://szlholdings.substack.com"),
        ("X handle",           "@szlholdings  (https://x.com/szlholdings)"),
        ("LinkedIn",           "https://linkedin.com/in/stephenlutar"),
        ("Best launch day",    "Tuesday"),
        ("Best launch hour",   "7:30 AM ET (LinkedIn) / 9:00 AM ET (X)"),
        ("Email send window",  "Tue–Wed, 6:45–7:30 AM ET"),
        ("Voice rails",        "Proof over pitch. No noise. Builder-operator. Trust-first."),
    ]
    for label, value in fields:
        add_label_value(doc, label, value)

    doc.add_paragraph()
    add_body(doc,
        "Files in this kit: SZL-Substack-Launch-Social-Kit.docx, build_docx.py",
        italic=True)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
